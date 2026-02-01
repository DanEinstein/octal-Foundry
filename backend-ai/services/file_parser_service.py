"""
File Parser Service
Extracts unit names from various file formats:
- TXT, CSV: Plain text parsing
- DOCX: Microsoft Word documents  
- XLSX: Excel spreadsheets
- Images (PNG, JPG): AI-powered OCR via Groq Vision
"""

import csv
import io
import re
import logging
import base64
from typing import List, Dict, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class FileParserService:
    """Parse various file formats to extract unit information"""
    
    SUPPORTED_EXTENSIONS = {
        'text': ['.txt', '.csv'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp']
    }
    
    def __init__(self, groq_api_key: str = None):
        self.groq_api_key = groq_api_key
    
    def parse_file(self, file_path: str, file_content: bytes = None) -> List[Dict]:
        """
        Parse a file and extract unit information
        
        Args:
            file_path: Path to file or filename (for extension detection)
            file_content: Raw file bytes (optional, will read from path if not provided)
            
        Returns:
            List of dicts with 'unit_code' and 'unit_name' keys
        """
        ext = Path(file_path).suffix.lower()
        
        # Read file if content not provided
        if file_content is None:
            with open(file_path, 'rb') as f:
                file_content = f.read()
        
        # Route to appropriate parser
        if ext in self.SUPPORTED_EXTENSIONS['text']:
            if ext == '.csv':
                return self._parse_csv(file_content)
            else:
                return self._parse_txt(file_content)
        elif ext in self.SUPPORTED_EXTENSIONS['word']:
            return self._parse_docx(file_content)
        elif ext in self.SUPPORTED_EXTENSIONS['excel']:
            return self._parse_xlsx(file_content)
        elif ext in self.SUPPORTED_EXTENSIONS['image']:
            return self._parse_image(file_content, ext)
        else:
            # Try as plain text
            logger.warning(f"Unknown extension {ext}, attempting plain text parse")
            return self._parse_txt(file_content)
    
    def _parse_txt(self, content: bytes) -> List[Dict]:
        """Parse plain text file (one unit per line)"""
        units = []
        try:
            text = content.decode('utf-8', errors='ignore')
            lines = text.strip().split('\n')
            
            for line in lines:
                line = line.strip()
                if not line or len(line) < 3:
                    continue
                    
                # Try to extract code and name
                unit = self._extract_unit_from_line(line)
                if unit:
                    units.append(unit)
                    
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            
        return units
    
    def _parse_csv(self, content: bytes) -> List[Dict]:
        """Parse CSV file"""
        units = []
        try:
            text = content.decode('utf-8', errors='ignore')
            reader = csv.reader(io.StringIO(text))
            
            # Skip header if it looks like one
            first_row = next(reader, None)
            if first_row and not any(kw in str(first_row).lower() for kw in ['code', 'unit', 'course', 'name']):
                # First row is data, not header
                if len(first_row) >= 2:
                    units.append({'unit_code': first_row[0], 'unit_name': first_row[1]})
                elif len(first_row) == 1:
                    units.append(self._extract_unit_from_line(first_row[0]))
            
            for row in reader:
                if not row or not row[0].strip():
                    continue
                if len(row) >= 2:
                    units.append({'unit_code': row[0].strip(), 'unit_name': row[1].strip()})
                elif len(row) == 1:
                    unit = self._extract_unit_from_line(row[0])
                    if unit:
                        units.append(unit)
                        
        except Exception as e:
            logger.error(f"Error parsing CSV: {e}")
            
        return units
    
    def _parse_docx(self, content: bytes) -> List[Dict]:
        """Parse Microsoft Word document"""
        units = []
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 3:
                    unit = self._extract_unit_from_line(text)
                    if unit:
                        units.append(unit)
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2 and cells[0] and cells[1]:
                        units.append({'unit_code': cells[0], 'unit_name': cells[1]})
                    elif len(cells) >= 1 and cells[0]:
                        unit = self._extract_unit_from_line(cells[0])
                        if unit:
                            units.append(unit)
                        
        except ImportError:
            logger.error("python-docx not installed")
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            
        return units
    
    def _parse_xlsx(self, content: bytes) -> List[Dict]:
        """Parse Excel spreadsheet"""
        units = []
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True)
            
            for sheet in wb.worksheets:
                skip_first = True
                for row in sheet.iter_rows(values_only=True):
                    if skip_first:
                        skip_first = False
                        # Check if this looks like a header
                        if row and any(str(cell).lower() in ['code', 'unit', 'name', 'course'] for cell in row if cell):
                            continue
                    
                    if not row or not row[0]:
                        continue
                        
                    if len(row) >= 2 and row[0] and row[1]:
                        units.append({'unit_code': str(row[0]).strip(), 'unit_name': str(row[1]).strip()})
                    elif row[0]:
                        unit = self._extract_unit_from_line(str(row[0]))
                        if unit:
                            units.append(unit)
                            
        except ImportError:
            logger.error("openpyxl not installed")
        except Exception as e:
            logger.error(f"Error parsing XLSX: {e}")
            
        return units
    
    def _parse_image(self, content: bytes, ext: str) -> List[Dict]:
        """Parse image using AI vision (OCR)"""
        units = []
        
        if not self.groq_api_key:
            logger.error("Groq API key required for image parsing")
            return units
            
        try:
            import requests
            
            # Encode image to base64
            base64_image = base64.b64encode(content).decode('utf-8')
            mime_type = f"image/{ext[1:]}" if ext != '.jpg' else 'image/jpeg'
            
            # Use Groq Vision API
            headers = {
                "Authorization": f"Bearer {self.groq_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": "llama-3.2-90b-vision-preview",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """Extract all course unit names from this image.
Return a JSON array where each item has:
- unit_code: The unit code if visible (e.g., "CIT 301"), or null
- unit_name: The full unit name

Return ONLY the JSON array, no other text. Example:
[{"unit_code": "CIT 301", "unit_name": "Machine Learning"}]"""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                "temperature": 0.3,
                "max_tokens": 2048
            }
            
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                text = result['choices'][0]['message']['content']
                
                # Extract JSON from response
                import json
                text = text.strip()
                if text.startswith('```'):
                    text = text.split('\n', 1)[1] if '\n' in text else text[3:]
                if text.endswith('```'):
                    text = text[:-3]
                
                start_idx = text.find('[')
                end_idx = text.rfind(']')
                if start_idx != -1 and end_idx != -1:
                    text = text[start_idx:end_idx+1]
                
                units = json.loads(text)
            else:
                logger.error(f"Vision API error: {response.status_code} - {response.text}")
                
        except Exception as e:
            logger.error(f"Error parsing image: {e}")
            
        return units
    
    def _extract_unit_from_line(self, line: str) -> Optional[Dict]:
        """Extract unit code and name from a line of text"""
        line = line.strip()
        if not line or len(line) < 3:
            return None
            
        # Skip common non-unit lines
        skip_keywords = ['semester', 'year', 'total', 'credit', 'hours', 'week']
        if any(kw in line.lower() for kw in skip_keywords):
            return None
        
        # Pattern: "ABC 123: Unit Name" or "ABC123 - Unit Name"
        pattern = r'^([A-Z]{2,5}\s*\d{2,4})\s*[:\-–—]\s*(.+)$'
        match = re.match(pattern, line, re.IGNORECASE)
        if match:
            return {'unit_code': match.group(1).strip().upper(), 'unit_name': match.group(2).strip()}
        
        # Pattern: "ABC 123 Unit Name" (code followed by name)
        pattern2 = r'^([A-Z]{2,5}\s*\d{2,4})\s+(.{5,})$'
        match2 = re.match(pattern2, line, re.IGNORECASE)
        if match2:
            return {'unit_code': match2.group(1).strip().upper(), 'unit_name': match2.group(2).strip()}
        
        # Just a unit name
        return {'unit_code': None, 'unit_name': line}
